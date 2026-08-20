from docx import Document
import sys

def extract_plan_completo(docx_path):
    try:
        doc = Document(docx_path)
        
        print("=== CONTENIDO COMPLETO DEL PLAN EXCLUSIVO ===")
        print()
        
        # Extraer todos los párrafos
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:
                print(f"[Párrafo {i}] {text}")
        
        print()
        print("=== TABLAS ===")
        print()
        
        # Extraer tablas
        for table_idx, table in enumerate(doc.tables, 1):
            print(f"Tabla {table_idx}:")
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                print(f"  Fila {row_idx}: {row_data}")
            print()
            
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    docx_path = "docs/plan_turistico_exclusivo/Plan salento filandia Y ocaso.docx"
    extract_plan_completo(docx_path)