from docx import Document
import sys
import os

def extract_plan_exclusivo(docx_path):
    try:
        if not os.path.exists(docx_path):
            print(f"Archivo no encontrado: {docx_path}")
            return
            
        doc = Document(docx_path)
        
        print("=== CONTENIDO DEL PLAN EXCLUSIVO ===")
        print()
        
        # Extraer todos los párrafos
        for i, para in enumerate(doc.paragraphs, 1):
            if para.text.strip():
                print(f"[Párrafo {i}] {para.text}")
        
        print()
        print("=== TABLAS ===")
        print()
        
        # Extraer tablas
        for table_idx, table in enumerate(doc.tables, 1):
            print(f"Tabla {table_idx}:")
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                print(f"  Fila {row_idx}: {row_data}")
            print()
            
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    docx_path = "docs/plan_turistico_exclusivo/Plan salento filandia Y ocaso.docx"
    extract_plan_exclusivo(docx_path)