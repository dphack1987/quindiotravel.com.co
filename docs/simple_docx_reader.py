from docx import Document
import sys

def simple_docx_reader(docx_path):
    try:
        doc = Document(docx_path)
        
        print("=== CONTENIDO DEL DOCUMENTO ===")
        print()
        
        # Extraer todos los párrafos
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            print(f"[{i}] ({len(text)} chars): {text}")
        
        print()
        print("=== TABLAS ===")
        print()
        
        # Extraer tablas
        print(f"Total de tablas: {len(doc.tables)}")
        for table_idx, table in enumerate(doc.tables, 1):
            print(f"TABLA {table_idx}:")
            print(f"  Filas: {len(table.rows)}")
            for row_idx, row in enumerate(table.rows, 1):
                row_data = [cell.text.strip() for cell in row.cells]
                print(f"  Fila {row_idx}: {row_data}")
            print()
            
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    docx_path = "plan_turistico_exclusivo/Plan salento filandia Y ocaso.docx"
    simple_docx_reader(docx_path)