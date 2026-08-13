from docx import Document
from pathlib import Path
import re

def limpiar_num(val):
    v = re.sub(r"[\s.,$COP]", "", val or "")
    return int(v) if v.isdigit() else val.strip() if val else ""

def extraer_todas_tablas(docx_path):
    doc = Document(str(docx_path))
    print(f"\n{'='*100}")
    print(f"📄 DOCUMENTO: {docx_path.name}")
    print(f"🔢 N° TABLAS TOTALES: {len(doc.tables)}")
    print(f"{'='*100}\n")
    for idx, tabla in enumerate(doc.tables):
        print(f"\n{'─'*90}")
        print(f"▶ TABLA #{idx} ({len(tabla.rows)} filas × {len(tabla.columns)} cols)")
        print(f"{'─'*90}")
        for f, fila in enumerate(tabla.rows):
            celdas = [celda.text.replace("\n", " | ").strip() for celda in fila.cells]
            # Normalizar precios: eliminar espacios intermedios $ ,
            celdas_print = []
            for c in celdas:
                c_clean = re.sub(r"\s+", " ", c)
                celdas_print.append(c_clean)
            print(f"  F{f:02d}:  {celdas_print}")

# ========================================================
# EJECUCIÓN
# ========================================================
base = Path(r"c:\Users\user\Documents\www.quindiotravel.com\docs")

doc1 = base / "informacion-de-precios" / "PORTAFOLIO PLANES NACIONALES 2026.docx"
extraer_todas_tablas(doc1)

doc2 = base / "promociones y precios para diciembre" / "planes especiales para diciembre con oferta max 30 cupos.docx"
extraer_todas_tablas(doc2)
