from docx import Document
import sys
import traceback
from pathlib import Path

def read_docx(file_path):
    try:
        # Convertir a Path y validar que existe
        file_path = Path(file_path)
        if not file_path.exists():
            print(f"Error: El archivo no existe: {file_path}")
            return False
        
        print("Opening file: " + str(file_path))
        doc = Document(str(file_path))
        
        print("=" * 80)
        print("FILE: " + str(file_path))
        print("=" * 80)
        print("Paragraphs: " + str(len(doc.paragraphs)))
        print("Tables: " + str(len(doc.tables)))
        print("=" * 80)
        print()
        
        # Read ALL paragraphs
        print("--- PARAGRAPH CONTENT ---")
        for i, para in enumerate(doc.paragraphs):
            text = para.text
            print(f"[Para {i+1}] (len: {len(text)}): '{text}'")
            if text.strip():
                print(f"  Content: {text}")
        
        # Read tables if they exist
        if doc.tables:
            print("\n" + "=" * 80)
            print("TABLES FOUND")
            print("=" * 80)
            for table_idx, table in enumerate(doc.tables):
                print(f"\n--- TABLE {table_idx + 1} ---")
                print(f"Dimensions: {len(table.rows)} rows x {len(table.columns)} columns")
                for row_idx, row in enumerate(table.rows):
                    row_text = [cell.text.strip() for cell in row.cells]
                    print(f"Row {row_idx + 1}: {' | '.join(row_text)}")
        else:
            print("\nNo tables found in document.")
        
        return True
    except Exception as e:
        print("Error reading file: " + str(e))
        traceback.print_exc()
        return False

if __name__ == "__main__":
    # Usar ruta relativa al directorio del script
    script_dir = Path(__file__).parent
    file_path = script_dir / "Pagina www.quindiotravel.com.co.docx"
    
    # Si no existe el archivo específico, mostrar ayuda
    if not file_path.exists():
        print("Uso: python read_docx.py [ruta_archivo]")
        print("Por defecto busca: Pagina www.quindiotravel.com.co.docx")
        print(f"Directorio actual: {script_dir}")
        
        # Intentar usar argumento de línea de comandos si se proporciona
        if len(sys.argv) > 1:
            file_path = Path(sys.argv[1])
        else:
            print("Por favor proporciona la ruta del archivo .docx como argumento")
            sys.exit(1)
    
    read_docx(file_path)